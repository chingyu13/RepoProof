from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path("/Users/chingyu/Projects/RepoProof")
SRC = ROOT / "docs" / "ethics" / "CARE - Substudy (Prospective) PIS v1b 2025-08-15.docx"
OUT = ROOT / "Done" / "RepoProof_COMP5339_Participant_Information_Statement_DRAFT.docx"

def set_text(paragraph, text):
    paragraph.clear()
    paragraph.add_run(text)

def remove_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)

def build():
    doc = Document(SRC)
    p = doc.paragraphs

    content = {
        0: "Participant Information Statement",
        1: "COMP5339 students — Semester 2, 2026",
        3: ("Research Study: Umbrella Ethics for prospective education research studies involving the evaluation "
            "of Units of Study delivered in Faculty of Engineering"),
        4: "Sub-study: RepoProof: Evaluating a personalised code-understanding quiz",
        6: "Ms Ching-Yu Chang (Responsible Researcher)",
        7: "School of Computer Science, Faculty of Engineering\nPhone: +61 4 3153 9704 | Email: ccha0131@uni.sydney.edu.au",

        10: "What is this study about?",
        12: ("This study evaluates RepoProof, a formative learning activity in COMP5339. RepoProof presents five "
             "multiple-choice questions based on a student’s submitted code. Teaching staff review the questions "
             "before release, and students receive immediate explanatory feedback. The study will examine whether "
             "this activity is feasible and useful, and will describe response, topic and completion-time patterns "
             "among students who consent to the research."),
        14: ("Taking part in the research is voluntary. The RepoProof quiz remains a COMP5339 teaching activity: "
             "students receive the 2% credit for completing it, not for answering correctly. Your research decision "
             "will not affect quiz access, feedback, marks or academic standing."),
        16: "Please read this sheet carefully and ask questions about anything you do not understand or would like to know more about.",
        18: "By giving consent, you confirm that you:",
        19: "have read and understood this Participant Information Statement;",
        20: "voluntarily agree to the research use described below; and",
        21: "agree to the collection, storage and use of your de-identified RepoProof research record as described.",
        23: "You may save or print a copy of this Participant Information Statement.",

        25: "Who is running the study?",
        27: "The study is being carried out by:",
        28: "Ching-Yu Chang, Student, School of Computer Science, Faculty of Engineering; and ?.",
        30: "No external funding has been identified.",

        32: "Who can take part in the study?",
        34: "Students enrolled in COMP5339 during Semester 2, 2026 may take part.",

        36: "What will the study involve for me?",
        37: ("As part of COMP5339 teaching, you will complete a five-question multiple-choice RepoProof quiz about "
             "your submitted code. It is expected to take about three minutes and provides immediate feedback."),
        39: "Optional consent for RepoProof research data",
        40: ("Before the quiz, you may choose whether your de-identified RepoProof research record can be retained "
             "for five years and used in individual and aggregate educational research analyses. The consent option "
             "is not pre-selected."),
        42: ("If you do not consent, you can still complete the quiz, receive feedback and receive the 2% completion "
             "credit. Your RepoProof record will not be retained or used for research, including aggregate research "
             "analysis. Operational data needed to deliver the quiz will be deleted approximately one month after "
             "the quiz deadline."),
        44: ("You may also complete a separate optional anonymous Qualtrics experience survey, which is expected "
             "to take approximately ? minutes. No interview, focus group, audio/video recording or additional "
             "access to your academic records is involved."),

        47: "Can I withdraw once I’ve started?",
        48: "Research participation is voluntary. You may decline without giving a reason and without disadvantage.",
        50: ("If you consent and later change your mind, contact ? at ? by ? (no later than one month after the quiz "
             "deadline). After the identifying linkage has been destroyed, or after information has been irreversibly "
             "de-identified and incorporated into non-identifiable results, it may no longer be possible to locate "
             "and remove your record."),
        58: ("For the anonymous Qualtrics survey, you may stop at any time before selecting Submit. Because the survey "
             "does not identify you, a response cannot be located or withdrawn after submission."),

        72: "Are there any risks or costs?",
        74: ("There is no financial cost. Possible risks are feeling pressure to participate and the small risk that "
             "learning data could be re-identified. Participation is therefore optional and separate from marks. "
             "RepoProof does not receive your SID; the identity linkage is held separately, identifying code content "
             "is removed before research retention, and access to research data is restricted."),

        75: "Are there any benefits?",
        77: ("There may be no direct research benefit to you beyond the feedback provided by the teaching activity. "
             "The findings may help improve RepoProof and future teaching in COMP5339."),

        79: "What will happen to information that is collected?",
        81: ("To run the quiz, RepoProof processes a RepoID and FolderID, generated and approved questions, answer "
             "options, selected responses, correctness, topic labels, timing information and code-derived content. "
             "RepoProof does not receive your SID."),
        83: ("For students who consent, the five-year research record may include quiz questions and responses, "
             "correctness, topic labels, duration information and approved de-identified information derived from "
             "code structure. Original source code, filenames, identifiers, strings, secrets and other identifying "
             "source text will not be retained in the five-year research record."),
        90: ("An independent CARE staff member will hold the SID–RepoID–FolderID linkage separately from the research "
             "data. Operational data will be kept on encrypted, access-controlled infrastructure in the AWS Sydney "
             "region and deleted approximately one month after the quiz deadline."),
        91: ("Approved de-identified research records will be transferred securely to the University Research Data "
             "Store using ?. Access will be limited to authorised research personnel. Records will be retained for "
             "five years and then permanently deleted under University procedures."),
        92: ("Results may be published or presented only in de-identified aggregate form. Reported groups will contain "
             "at least five consenting participants and will be checked to reduce disclosure risk. The participation "
             "rate will be reported, and findings will not be described as necessarily representative of the full "
             "COMP5339 cohort."),
        93: ("With your consent, the de-identified record may also be used in future educational research only where "
             "the required ethics approval is in place."),

        95: "Will I be told the results of the study?",
        97: "A short plain-language summary of the overall results will be made available through ? after the study.",

        99: "What if I would like further information?",
        101: ("For further information or questions about the study, contact Ching-Yu Chang, Student, School of "
              "Computer Science, at ccha0131@uni.sydney.edu.au or +61 4 3153 9704."),

        104: "What if I have a complaint or any concerns?",
        106: ("The ethical aspects of this study are being considered under the University of Sydney Faculty of "
              "Engineering CARE umbrella approval, HREC 2025/HE001132, sub-study ID ?. Research use will not begin "
              "until the required sub-study approval has been confirmed."),
        108: ("If you are concerned about the conduct of the study or wish to make a complaint to someone independent "
              "of the study, please contact:"),
        110: "Human Ethics Manager",
        111: "human.ethics@sydney.edu.au",
        112: "+61 2 8627 8176",
        117: "This information sheet is for you to keep.",
    }

    for index, text in content.items():
        set_text(p[index], text)
    keep = set(content)
    for index, paragraph in reversed(list(enumerate(doc.paragraphs))):
        if index not in keep:
            remove_paragraph(paragraph)

    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.0
        for run in paragraph.runs:
            if run.font.size is None or run.font.size.pt > 10.5:
                run.font.size = Pt(10.5)

    for section in doc.sections:
        footer = section.footer
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.text = "DRAFT — fields marked ? require confirmation before submission"
        for run in fp.runs:
            run.font.size = Pt(8)

    doc.save(OUT)
    print(OUT)

if __name__ == "__main__":
    build()
