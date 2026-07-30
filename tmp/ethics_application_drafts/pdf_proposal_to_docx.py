from pathlib import Path
from pypdf import PdfReader
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path("/Users/chingyu/Projects/RepoProof")
PDF = ROOT / "Done" / "RepoProof_COMP5339_CARE_Umbrella_SubStudy_Proposal_Response_DRAFT.pdf"
OUT = ROOT / "Done" / "RepoProof_COMP5339_CARE_Umbrella_SubStudy_Proposal_Response_DRAFT.docx"

TITLE = "CARE Umbrella Ethics — Sub-Study Proposal Response Draft"
SUBTITLE = "RepoProof: Evaluating a personalised code-understanding quiz in COMP5339"
NOTE = "Prepared for entry into the CARE Microsoft Form. This is not an approval. All fields marked ? must be confirmed before submission."

HEADINGS = [
    "Application type",
    "Sub-study title",
    "Chief Investigator",
    "Other research team members",
    "Additional notes",
    "Prospective or retrospective?",
    "Aims and research questions",
    "Background and rationale (under 500 words)",
    "Study duration",
    "Participants",
    "Planned sample size",
    "Study location",
    "Recruitment and consent",
    "Data collected",
    "Methodology and analysis",
    "Confidentiality and privacy",
    "Data protection, retention and disposal",
    "Risks and mitigation",
    "Expected significance",
    "Attachments",
    "Acknowledgements / declarations",
    "Outstanding confirmations before submission",
]

BLUE = RGBColor(22, 59, 100)
MUTED = RGBColor(105, 105, 105)

def set_font(run, size, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color

def extract_sections():
    raw_lines = []
    for page in PdfReader(str(PDF)).pages:
        for line in (page.extract_text() or "").splitlines():
            s = line.strip()
            if not s or s.startswith("DRAFT — fields marked") or s.startswith("Page "):
                continue
            raw_lines.append(s)

    start = raw_lines.index(HEADINGS[0])
    lines = raw_lines[start:]
    sections = []
    current_heading = None
    current_lines = []
    for line in lines:
        if line in HEADINGS:
            if current_heading is not None:
                sections.append((current_heading, " ".join(current_lines)))
            current_heading = line
            current_lines = []
        else:
            current_lines.append(line)
    if current_heading is not None:
        sections.append((current_heading, " ".join(current_lines)))
    return sections

def add_page_field(paragraph):
    paragraph.add_run("Page ")
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)

def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(18)
    section.bottom_margin = Mm(18)
    section.left_margin = Mm(21)
    section.right_margin = Mm(21)
    section.header_distance = Mm(10)
    section.footer_distance = Mm(10)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    h1._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    h1.font.size = Pt(13)
    h1.font.bold = True
    h1.font.color.rgb = BLUE
    h1.paragraph_format.space_before = Pt(10)
    h1.paragraph_format.space_after = Pt(3)
    h1.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(header.add_run("CARE Umbrella Ethics | COMP5339 RepoProof"), 8.5, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(footer.add_run("DRAFT — fields marked ? require confirmation before submission  |  "), 8, color=MUTED)
    add_page_field(footer)
    for r in footer.runs:
        set_font(r, 8, color=MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(7)
    set_font(p.add_run(TITLE), 20, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(5)
    set_font(p.add_run(SUBTITLE), 15, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(15)
    set_font(p.add_run(NOTE), 9, color=MUTED)

    for heading, answer in extract_sections():
        doc.add_paragraph(heading, style="Heading 1")
        p = doc.add_paragraph(answer)
        p.paragraph_format.keep_together = False
        p.paragraph_format.widow_control = True

    doc.core_properties.title = TITLE
    doc.core_properties.subject = SUBTITLE
    doc.core_properties.author = "?"
    doc.save(OUT)
    print(OUT)

if __name__ == "__main__":
    build()
