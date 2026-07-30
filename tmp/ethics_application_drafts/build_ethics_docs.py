from pathlib import Path
from copy import deepcopy
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle, KeepTogether

ROOT = Path("/Users/chingyu/Projects/RepoProof")
SRC = ROOT / "docs" / "ethics"
OUT = ROOT / "Done"
OUT.mkdir(exist_ok=True)

TITLE = "RepoProof: Evaluating a personalised code-understanding quiz in COMP5339"
CI = "? (Chief Investigator and Responsible Researcher)"
CONTACT = "School of Computer Science, Faculty of Engineering\nPhone: ? | Email: ?"

def set_para(p, text, bold=False):
    p.clear()
    r = p.add_run(text)
    r.bold = bold
    return p

def remove_para(p):
    el = p._element
    el.getparent().remove(el)

def keep_only(doc, indices):
    for i, p in reversed(list(enumerate(doc.paragraphs))):
        if i not in indices:
            remove_para(p)

def add_bullet_after(paragraph, text):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    p = paragraph._parent.add_paragraph()
    p._p.getparent().remove(p._p)
    new_p.addnext(p._p)
    p.style = "List Paragraph"
    p.add_run(text)
    return p

def footer_draft(doc):
    for section in doc.sections:
        f = section.footer
        p = f.paragraphs[0] if f.paragraphs else f.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.text = "DRAFT — fields marked ? require confirmation before submission"
        for r in p.runs:
            r.font.size = Pt(8)

def build_pis():
    src = SRC / "CARE - Substudy (Prospective) PIS v1b 2025-08-15.docx"
    doc = Document(src)
    P = doc.paragraphs
    repl = {
        0: "Participant Information Statement",
        1: "COMP5339 students — Semester 2, 2026",
        3: "Research Study: Umbrella Ethics for prospective education research studies involving the evaluation of Units of Study delivered in Faculty of Engineering",
        4: f"Sub-study: {TITLE}",
        6: CI,
        7: CONTACT,
        10: "What is this study about?",
        12: ("This study evaluates RepoProof, a formative learning activity in COMP5339. RepoProof creates five "
             "multiple-choice questions about a student’s submitted code. Questions are reviewed by teaching staff "
             "before release, and students receive immediate explanatory feedback. The research asks whether this "
             "approach is feasible and acceptable, what topic and completion patterns occur, and whether patterns "
             "change across semesters or following teaching changes. It is not an authorship, plagiarism, misconduct "
             "or AI-detection tool."),
        17: ("Taking part in the research is voluntary. Completing the quiz earns the advertised 2% course credit; "
             "credit is for completion only, not for correct answers. Your research choices do not affect that credit, "
             "your marks, teaching access, or relationship with the University."),
        19: "Please read this sheet carefully and ask questions about anything you do not understand or want to know more about.",
        21: "The study has one explicit opt-in pathway for RepoProof research data:",
        22: ("Your individual de-identified RepoProof research record will be retained for five years and included in "
             "individual and aggregate analyses only if you explicitly opt in."),
        23: ("If you do not opt in, you can still complete the quiz and receive the 2% completion credit. Your data will "
             "be used only to operate the teaching activity and will be deleted about one month after the quiz deadline."),
        24: ("The anonymous Qualtrics experience survey is separately optional. It is included in research only if you submit it."),
        26: "You may save or print a copy of this Participant Information Statement.",
        28: "Who is running the study?",
        30: "The study is being carried out by:",
        31: f"{CI}, {CONTACT.replace(chr(10), ', ')}; other research team members: ?.",
        33: "No external funding has been identified.",
        36: "Who can take part in the study?",
        38: ("Students enrolled in COMP5339 during Semester 2, 2026 are eligible. Students complete the teaching activity "
             "whether or not they allow research use of their data. No additional exclusion criteria are proposed."),
        40: "What will the study involve for me?",
        43: ("As part of COMP5339 teaching, you will complete a five-question multiple-choice RepoProof quiz about your "
             "submitted code. It is expected to take about three minutes and provides immediate feedback. Before the "
             "quiz, you may choose whether to allow an individual de-identified research record to be kept for five "
             "years. You may also complete an optional anonymous Qualtrics experience survey. No interview, focus "
             "group, audio/video recording, or additional access to your academic record is involved."),
        56: "Explicit opt-in for all RepoProof research use",
        59: ("If you opt in, your de-identified RepoProof record may be retained for five years and used in both individual "
             "and aggregate analyses. Aggregate analysis may include topic-response patterns, earlier-versus-later "
             "completion groups, duration bands, and comparisons across semesters or teaching changes. Reports will "
             "describe consenting participants only, report participation rates, use groups of at least five, and avoid "
             "complementary-cell disclosure."),
        61: ("If you do not opt in, your quiz record will not be retained or used for research, including aggregate "
             "research analysis. You may still complete the quiz normally and receive the 2% completion credit."),
        63: ("The research opt-in choice is displayed before the quiz and is not pre-selected. Choosing not to opt in "
             "has no effect on quiz access, feedback, marks or academic standing."),
        67: "Can I withdraw once I’ve started?",
        68: "Research participation is completely voluntary.",
        70: "Your decision will not affect your current or future relationship with the researchers or anyone else at the University of Sydney.",
        72: ("If you opt in and later change your mind, contact the independent CARE staff member at ? by ? (no later "
             "than one month after the quiz deadline). After the linkage is destroyed or the record has been irreversibly "
             "de-identified and incorporated into non-identifiable results, it may no longer be possible to locate and remove it."),
        78: ("For the anonymous Qualtrics experience survey, you may withdraw any time before pressing Submit. Once "
             "submitted, the survey cannot be linked back to you and therefore cannot be withdrawn."),
        92: "Are there any risks or costs?",
        94: ("The main risks are discomfort about research use of learning data and a low risk of re-identification. "
             "Controls include separation of identity mapping from RepoProof, encrypted storage, access limited to "
             "authorised personnel, removal of identifiers and identifying content, minimum reported cell sizes of "
             "five, and no effect of research choices or answer correctness on marks. There is no financial cost."),
        95: "Are there any benefits?",
        97: ("There may be no direct benefit to you beyond the formative quiz feedback. Findings may improve teaching and "
             "assessment design in COMP5339 and inform responsible use of personalised code-understanding quizzes."),
        99: "What will happen to information that is collected?",
        101: ("RepoProof receives a RepoID and FolderID, not your SID. Operational data may include quiz questions, answer "
              "options, selected responses, correctness, topic labels, timing/duration information and code-derived "
              "content needed to run the quiz. Source submissions and operational data are manually deleted one month "
              "after the quiz deadline."),
        103: ("The independent CARE staff member holds the SID–RepoID–FolderID linkage separately. Before research "
              "retention, identifiers and identifying code content—including filenames, identifiers, strings, secrets "
              "and source text—will be removed or transformed. The research team analyses only de-identified data."),
        110: ("During the teaching activity, RepoProof is hosted on an AWS EC2 instance in the Sydney region with an "
              "encrypted EBS volume (AWS-managed key alias/aws/ebs). Production access is currently limited to the "
              "student developer. There is no routine application backup. Any temporary migration snapshot will be "
              "deleted after verification. A data breach will be contained and reported under University of Sydney "
              "privacy, ICT and research-governance procedures."),
        111: ("Approved de-identified research data will be securely transferred to the University Research Data Store "
              "using ?. Access will be restricted to authorised research personnel. It will be retained for five years "
              "and then permanently deleted under University of Sydney Archives and records-management procedures."),
        112: ("Only de-identified aggregate results will be reported. Published tables and figures will use groups of at "
              "least five and will be checked to avoid complementary-cell or combination disclosure."),
        113: ("If you explicitly opt in, the de-identified record may be used in individual and aggregate analyses for "
              "this study and future educational research that has appropriate ethics approval. If you do not opt in, "
              "your RepoProof record will not be retained or included in research analyses."),
        115: "Will I be told the results of the study?",
        117: ("A short plain-language summary of overall results will be made available through ? after the study. "
              "Do not provide contact details in the anonymous survey solely to receive feedback."),
        119: "What if I would like further information?",
        121: f"Please contact {CI}: phone ?, email ?.",
        124: "What if I have a complaint or any concerns?",
        126: ("The ethical aspects of this study are being considered under the University of Sydney Faculty of "
              "Engineering CARE umbrella approval, HREC 2025/HE001132, sub-study ID ?. Research use will not begin "
              "until the required sub-study approval is confirmed."),
        128: ("If you are concerned about the conduct of the study or wish to complain to someone independent of the "
              "study, please contact the University:"),
        130: "Human Ethics Manager",
        131: "human.ethics@sydney.edu.au",
        132: "+61 2 8627 8176",
        137: "This information sheet is for you to keep.",
    }
    keep = set(repl)
    for i, text in repl.items():
        set_para(P[i], text)
    keep_only(doc, keep)
    footer_draft(doc)
    out = OUT / "RepoProof_COMP5339_Participant_Information_Statement_DRAFT.docx"
    doc.save(out)
    return out

def build_consent():
    src = SRC / "CARE - Substudy (Prospective) Consent v1b 2025-08-15.docx"
    doc = Document(src)
    P = doc.paragraphs
    repl = {
        0: "Participant E-Consent Form",
        1: "COMP5339 students — optional five-year RepoProof research record",
        3: "Research Study: Umbrella Ethics for prospective education research studies involving the evaluation of Units of Study delivered in University of Sydney Faculty of Engineering",
        5: f"Sub-study: {TITLE}",
        7: CI,
        8: CONTACT,
        12: ("I voluntarily consent to my de-identified RepoProof research record being retained for five years and "
             "used in individual and aggregate research analyses. This choice is separate from completing the quiz "
             "and from the optional anonymous survey. In giving consent, I confirm that:"),
        13: "I have read and understood the Participant Information Statement and have had an opportunity to ask questions.",
        14: ("I understand that the study evaluates the feasibility, acceptability and learning-data patterns of an "
             "instructor-reviewed personalised code-understanding quiz in COMP5339."),
        15: "I understand the possible risks and benefits described in the Participant Information Statement.",
        16: ("I will complete the normal five-question RepoProof teaching quiz. Giving or refusing this research consent "
             "does not change the activity, my 2% completion credit, or any other mark."),
        18: ("I consent to my de-identified quiz and code-derived research record being retained for five years and used "
             "in individual and aggregate analyses for this study and future educational research only where appropriate "
             "ethics approval is in place."),
        19: "I understand that participation is completely voluntary.",
        20: "I understand that my choice will not affect my marks, teaching access, or relationship with the research team or University.",
        21: ("I understand that I may request withdrawal through the independent CARE staff member at ? before the "
             "operational linkage is destroyed one month after the quiz deadline. After irreversible de-identification "
             "and inclusion in non-identifiable results, removal may no longer be possible."),
        22: ("I understand that RepoProof does not receive my SID; the linkage is held separately by independent CARE "
             "staff. Identifiers and identifying code content will be removed or transformed before five-year retention."),
        23: ("I understand that analyses include consenting participants only, participation rates will be reported, "
             "and results will not be described as necessarily representative of the full COMP5339 cohort."),
        27: "Please select one option:",
        29: ("☐ YES — I consent to five-year retention and individual and aggregate research use of my de-identified "
             "RepoProof record."),
        31: ("☐ NO — do not retain or use my RepoProof record for research, including aggregate research analysis. "
             "This does not affect quiz access, feedback, credit or marks."),
        41: "I understand that I may save or request a copy of this e-consent record.",
    }
    for i, text in repl.items():
        set_para(P[i], text)
    keep_only(doc, set(repl))
    for p in doc.paragraphs:
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0
        for r in p.runs:
            r.font.size = Pt(9.5)
    for i in (27, 29, 31, 41):
        if i < len(P) and P[i]._element.getparent() is not None:
            P[i].style = doc.styles["Normal"]
            ppr = P[i]._p.get_or_add_pPr()
            numpr = ppr.find(qn("w:numPr"))
            if numpr is not None:
                ppr.remove(numpr)
    if doc.tables:
        t0 = doc.tables[0]
        t0.cell(0,0).text = "RepoID (do not enter your SID)"
        t0.cell(0,1).text = "?"
        for t in doc.tables[1:]:
            t._element.getparent().remove(t._element)
    footer_draft(doc)
    out = OUT / "RepoProof_COMP5339_Individual_Record_EConsent_DRAFT.docx"
    doc.save(out)
    return out

def build_recruitment(kind):
    optout = kind == "optout"
    name = "CARE - Substudy (Prospective) Recruitment - optout v1b 2025-08-15.docx" if optout else "CARE - Substudy (Prospective) Recruitment - optin v1b 2025-08-15.docx"
    doc = Document(SRC / name)
    P = doc.paragraphs
    if optout:
        repl = {
            1: "Educational Research — Umbrella Ethics",
            3: "Announcement for study participation\nOPT-OUT",
            4: "(Routinely collected quiz data — cohort aggregate analysis)",
            5: "?",
            8: "Announcement schedule — before the COMP5339 RepoProof quiz; reminder before the quiz deadline ?",
            9: f"Announcement header — {TITLE}: opt-out of cohort aggregate research analysis",
            11: "Dear COMP5339 students,",
            12: ("We invite your routinely collected RepoProof quiz data to contribute to research evaluating an "
                 "instructor-reviewed personalised code-understanding quiz and improving teaching in COMP5339."),
            13: "The attached Participant Information Statement explains the study, privacy controls and your choices.",
            14: ("The proposed CARE-approved opt-out pathway applies only to de-identified cohort aggregate analysis of "
                 "routine quiz data. It does not authorise five-year retention of your individual record."),
            15: ("If you are willing for your quiz data to contribute to grouped analysis, you do not need to do anything. "
                 "There is no additional activity or time commitment."),
            17: f"You will be included in the cohort aggregate evaluation of {TITLE} unless you opt out by the deadline.",
            19: ("To opt out, an independent CARE staff member will confidentially process your request and exclude your "
                 "data using a separately held SID–RepoID–FolderID mapping. The lecturer and researchers will not be "
                 "told who opted out. You must still complete the teaching activity; your choice has no effect on the "
                 "2% quiz-completion credit or any other mark."),
            21: "Confidential Qualtrics revocation form: ?. Participation in research use is voluntary and there are no consequences for opting out.",
            22: "Final opportunity to opt out: ? (the quiz deadline).",
            23: "Questions: ? (Chief Investigator), ?.",
            24: "HREC approval number 2025/HE001132 — sub-study ID ?.",
            25: "Kind regards,",
            26: "Research team: ?",
        }
    else:
        repl = {
            1: "Educational Research — Umbrella Ethics",
            3: "Announcement for study participation\nOPT-IN",
            4: "(Optional five-year RepoProof research record and anonymous experience survey)",
            5: "?",
            8: "Announcement schedule — before the COMP5339 RepoProof quiz; reminder before the quiz deadline ?",
            9: f"Announcement header — {TITLE}: optional research participation",
            11: "Dear COMP5339 students,",
            12: ("We invite you to choose whether your de-identified RepoProof quiz and code-derived research record may "
                 "be retained for five years and used in individual and aggregate educational research analyses. We also "
                 "invite you to complete a separate anonymous experience survey."),
            13: "The attached Participant Information Statement explains the study, data handling, risks and your choices.",
            14: ("Review the PIS and e-consent at ?. Select YES only if you consent to five-year retention and all "
                 "research use of your de-identified RepoProof record, including aggregate analysis. If you select NO or "
                 "do not opt in, your record will not be retained or used for research. This has no effect on the 2% "
                 "quiz-completion credit or any other mark. The optional anonymous survey is at ?."),
            15: ("The normal RepoProof quiz contains five multiple-choice questions and is expected to take about three "
                 "minutes. The anonymous experience survey takes approximately ? minutes. No interview, focus group or recording is involved."),
            16: "Final opportunity to opt in / complete the anonymous survey: ?.",
            17: "Questions: ? (Chief Investigator), ?.",
            18: "HREC approval number 2025/HE001132 — sub-study ID ?.",
            19: "Kind regards,",
            20: "Research team: ?",
        }
    for i, text in repl.items():
        set_para(P[i], text)
    keep_only(doc, set(repl))
    footer_draft(doc)
    suffix = "OptOut_Cohort_Aggregates" if optout else "OptIn_All_RepoProof_Research_Use_and_Survey"
    out = OUT / f"RepoProof_COMP5339_Recruitment_{suffix}_DRAFT.docx"
    doc.save(out)
    return out

def page_number(canvas, doc):
    canvas.saveState()
    canvas.setFont("Helvetica", 8)
    canvas.setFillColor(colors.grey)
    canvas.drawString(20*mm, 12*mm, "DRAFT — fields marked ? require confirmation before submission")
    canvas.drawRightString(190*mm, 12*mm, f"Page {doc.page}")
    canvas.restoreState()

def build_umbrella_pdf():
    path = OUT / "RepoProof_COMP5339_CARE_Umbrella_SubStudy_Proposal_Response_DRAFT.pdf"
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="Title2", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=18, leading=22, textColor=colors.HexColor("#163B64"), alignment=TA_CENTER, spaceAfter=10))
    styles.add(ParagraphStyle(name="Q", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=11, leading=14, textColor=colors.HexColor("#163B64"), spaceBefore=7, spaceAfter=3))
    styles.add(ParagraphStyle(name="A", parent=styles["BodyText"], fontSize=9.3, leading=13, spaceAfter=5))
    styles.add(ParagraphStyle(name="Small", parent=styles["BodyText"], fontSize=8, leading=10, textColor=colors.grey))
    doc = SimpleDocTemplate(str(path), pagesize=A4, rightMargin=18*mm, leftMargin=18*mm, topMargin=18*mm, bottomMargin=20*mm, title=TITLE)
    story = [
        Paragraph("CARE Umbrella Ethics — Sub-Study Proposal Response Draft", styles["Title2"]),
        Paragraph(TITLE, styles["Heading1"]),
        Paragraph("Prepared for entry into the CARE Microsoft Form. This is not an approval. All fields marked ? must be confirmed before submission.", styles["Small"]),
        Spacer(1, 6),
    ]
    qa = [
        ("Application type", "New sub-study proposal under HREC 2025/HE001132."),
        ("Sub-study title", TITLE),
        ("Chief Investigator", "?"),
        ("Other research team members", "?"),
        ("Additional notes", ("Prospective educational evaluation in COMP5339, Semester 2 2026. CARE Working Party approval "
                              "and confirmation that COMP5339 is covered by the CARE Programs and Units list are pending. "
                              "No research analysis or five-year research transfer will begin before approval.")),
        ("Prospective or retrospective?", "Prospective."),
        ("Aims and research questions", ("Aim: evaluate the feasibility, acceptability and educational utility of RepoProof, an "
          "instructor-reviewed personalised code-understanding quiz. Questions: (1) Can five grounded MCQs and immediate "
          "feedback be delivered reliably within the COMP5339 workflow? (2) What topic-response, completion and duration "
          "distributions occur? (3) Do aggregate patterns differ between earlier and later completers, semesters, or following "
          "documented teaching changes? (4) What do students report about usefulness, clarity and perceived fairness in an "
          "optional anonymous survey? The study does not assess authorship, plagiarism, misconduct or AI use.")),
        ("Background and rationale (under 500 words)", ("Formative assessment can help learners regulate learning when it "
          "provides timely information about current understanding and next steps. Retrieval practice also supports durable "
          "learning. In programming education, generic questions may not address the specific structures and misconceptions "
          "visible in a student’s own submission. RepoProof therefore generates five multiple-choice questions grounded in "
          "submitted code; teaching staff approve questions before release and explanatory feedback follows each response. "
          "The quiz is a low-stakes teaching activity: 2% credit is awarded for completion, not correctness. Research is needed "
          "to determine whether the workflow is feasible and acceptable and whether grouped response and duration patterns "
          "provide useful signals for teaching improvement without retaining raw submissions. All RepoProof research use "
          "requires explicit opt-in: only consenting participants’ de-identified records are retained for five years and used "
          "in individual or aggregate analyses. The anonymous experience survey is separately optional. Identity linkage is "
          "held by independent CARE staff, and the research "
          "team receives only de-identified data. References: Nicol, D.J. & Macfarlane-Dick, D. (2006), Studies in Higher "
          "Education 31(2), 199–218; Roediger, H.L. & Karpicke, J.D. (2006), Psychological Science 17(3), 249–255; "
          "NHMRC, ARC & Universities Australia, National Statement on Ethical Conduct in Human Research (2025); University "
          "of Sydney Research Data Management Policy 2014 and Procedures 2015.")),
        ("Study duration", "Semester 2, 2026 to ?. Data collection date(s): ?. Five-year retention ends: ?."),
        ("Participants", ("Students enrolled in COMP5339 in Semester 2, 2026. Students undertake the quiz as a normal learning "
                          "activity regardless of research participation. No additional exclusion criteria are proposed.")),
        ("Planned sample size", ("All enrolled COMP5339 students will be invited; the research sample will comprise only "
                                 "students who explicitly opt in. Expected enrolment n = ?; expected consent rate and resulting "
                                 "research sample n = ?. Optional survey response count is unknown.")),
        ("Study location", ("Online. RepoProof application and SQLite database on AWS EC2 in ap-southeast-2 (Sydney); optional "
                            "survey and e-consent in University-approved Qualtrics and/or the approved RepoProof consent page; approved five-year research data "
                            "in the University Research Data Store.")),
        ("Recruitment and consent", ("The PIS and opt-in announcement will be provided before the quiz. All RepoProof research "
          "use requires explicit opt-in through an unselected consent control before the quiz. Only consenting participants’ "
          "de-identified records are retained for five years and included in individual or aggregate research analyses. "
          "Students who do not opt in still complete the quiz and receive the 2% completion credit; their operational data "
          "are deleted about one month after the quiz deadline and are not used for research. The anonymous experience survey "
          "is separately voluntary and opt-in.")),
        ("Data collected", ("Operational: RepoID, FolderID, generated/approved question text and options, selected response, "
          "correctness, explanatory feedback, topic labels, timestamps/duration and code-derived content required to operate "
          "the quiz. RepoProof does not receive SID. Optional survey: anonymous ratings and free-text responses. The independent "
          "CARE staff member alone holds the SID–RepoID–FolderID mapping. No interviews, recordings or direct academic-record access.")),
        ("Methodology and analysis", ("Among consenting participants, descriptive analysis of completion, topic-response and duration distributions; comparisons "
          "between predefined earlier/later completion groups and, if repeated in later semesters, across semesters or documented "
          "teaching changes. Analyses will be exploratory unless a statistical analysis plan is approved before access. Only "
          "de-identified data will be analysed. Consent participation rates will be reported, and findings will not be "
          "interpreted as necessarily representative of the full COMP5339 cohort. Aggregate outputs require n ≥ 5 and will be checked for complementary-cell "
          "disclosure. Optional anonymous survey items will be summarised descriptively; free text will be reviewed and redacted "
          "for inadvertent identifiers before quotation. No automated misconduct or authorship inference will be made.")),
        ("Confidentiality and privacy", ("SID never reaches RepoProof. Independent CARE staff holds the linkage separately and "
          "processes consent-related withdrawal requests. Before research retention, the pipeline removes or transforms filenames, "
          "identifiers, strings, secrets and source text. Research staff receive de-identified records only. The lecturer will not "
          "be told who consented or declined. Publications contain only de-identified aggregates.")),
        ("Data protection, retention and disposal", ("During delivery, the application runs on AWS EC2 in Sydney with encrypted "
          "EBS using alias/aws/ebs; access is currently limited to the student developer. Individual accounts and least-privilege "
          "roles will be used if access expands; AWS IAM is limited to infrastructure administrators. No routine application backup "
          "is kept. Any temporary migration snapshot will be deleted after verification. Operational submissions and linkage are "
          "manually deleted one month after the quiz deadline, with a deletion log. Approved de-identified research data will be "
          "transferred to the University Research Data Store using ?, restricted to authorised researchers, retained for five years, "
          "then deleted under University Archives procedures. Incidents will be contained and reported under University privacy, "
          "ICT and research-governance procedures.")),
        ("Risks and mitigation", ("Low risk. Potential concerns include perceived coercion because researchers are connected to "
          "teaching, discomfort about code-based questions, re-identification, and selection bias from consent-based participation. "
          "Mitigations: completion-only credit; research choices separated from marks; explicit opt-in for every RepoProof "
          "research use; reporting of consent participation rates and limits on representativeness; no SID in "
          "RepoProof; content de-identification; encrypted and access-controlled storage; operational deletion; n ≥ 5 reporting; "
          "and clear notice that RepoProof is not a misconduct or AI-detection tool.")),
        ("Expected significance", ("Evidence about whether instructor-reviewed personalised code questions can provide useful "
          "formative feedback while maintaining clear consent boundaries and data minimisation. Findings may improve COMP5339 "
          "teaching and inform responsible design of similar educational tools.")),
        ("Attachments", ("1. Participant Information Statement draft. 2. RepoProof research e-consent draft. 3. Recruitment "
          "opt-in announcement draft. 4. Current RepoProof product/ethics design deck. 5. De-identification and deletion "
          "implementation evidence, to be attached before research commences: ?.")),
        ("Acknowledgements / declarations", ("The team will comply with HREC 2025/HE001132, CARE sub-study conditions, the "
          "National Statement, University research-data requirements, annual reporting and incident/adverse-event reporting. "
          "Chief Investigator acknowledgement: ?. Research team acknowledgement(s): ?. Date: ?.")),
        ("Outstanding confirmations before submission", ("CI and team details; independent CARE staff member; COMP5339 coverage; "
          "sub-study ID; sample size; dates/deadline; Qualtrics URLs; survey duration; result-feedback channel; secure RDS transfer "
          "method; completed de-identification/deletion testing; removal of temporary unencrypted snapshot.")),
    ]
    for q, a in qa:
        story.append(KeepTogether([Paragraph(q, styles["Q"]), Paragraph(a, styles["A"])]))
    doc.build(story, onFirstPage=page_number, onLaterPages=page_number)
    return path

if __name__ == "__main__":
    outputs = [
        build_umbrella_pdf(),
        build_pis(),
        build_consent(),
        build_recruitment("optin"),
    ]
    print("\n".join(str(p) for p in outputs))
