from pathlib import Path
from docx import Document

ROOT = Path("/Users/chingyu/Projects/RepoProof")
DONE = ROOT / "Done"
OLD = ROOT / "tmp" / "ethics_application_drafts" / "obsolete_optout"

def replace_exact(doc, mapping, label):
    found = set()
    for p in doc.paragraphs:
        if p.text in mapping:
            old = p.text
            p.text = mapping[old]
            found.add(old)
    missing = set(mapping) - found
    if missing:
        raise RuntimeError(f"{label}: missing paragraphs: {sorted(missing)}")

def update_pis():
    path = DONE / "RepoProof_COMP5339_Participant_Information_Statement_DRAFT.docx"
    doc = Document(path)
    mapping = {
        "The study has three separate participation pathways:":
            "The study has one explicit opt-in pathway for RepoProof research data:",
        "1. Cohort aggregate analytics: routinely collected quiz data may be used in grouped analysis unless you opt out by the deadline using the confidential Qualtrics form.":
            "Your individual de-identified RepoProof research record will be retained for five years and included in individual and aggregate analyses only if you explicitly opt in.",
        "2. Individual five-year research record: your de-identified record is retained only if you explicitly opt in.":
            "If you do not opt in, you can still complete the quiz and receive the 2% completion credit. Your data will be used only to operate the teaching activity and will be deleted about one month after the quiz deadline.",
        "3. Anonymous experience survey: the optional Qualtrics survey is separate and is included only if you submit it.":
            "The anonymous Qualtrics experience survey is separately optional. It is included in research only if you submit it.",
        "Opt-out consent for routinely collected cohort aggregate data":
            "Explicit opt-in for all RepoProof research use",
        "If you are willing for your routinely collected quiz data to contribute to cohort aggregate analysis, you do not need to do anything further. This analysis creates distributions such as topic response patterns, earlier-versus-later completion groups, duration bands, and comparisons across semesters or teaching changes. Reports use groups of at least five and avoid complementary-cell disclosure.":
            "If you opt in, your de-identified RepoProof record may be retained for five years and used in both individual and aggregate analyses. Aggregate analysis may include topic-response patterns, earlier-versus-later completion groups, duration bands, and comparisons across semesters or teaching changes. Reports will describe consenting participants only, report participation rates, use groups of at least five, and avoid complementary-cell disclosure.",
        "You will be included in the cohort aggregate evaluation of RepoProof: Evaluating a personalised code-understanding quiz in COMP5339 unless you opt out by ? (the quiz deadline).":
            "If you do not opt in, your quiz record will not be retained or used for research, including aggregate research analysis. You may still complete the quiz normally and receive the 2% completion credit.",
        "To opt out, submit the confidential Qualtrics revocation form at ?. An independent CARE staff member, who is not part of the teaching/research team, will use the separately held SID–RepoID–FolderID mapping to exclude your data. The researchers and lecturer will not be told who opted out. You must still complete required teaching activities, and opting out has no effect on the 2% completion credit or any mark.":
            "The research opt-in choice is displayed before the quiz and is not pre-selected. Choosing not to opt in has no effect on quiz access, feedback, marks or academic standing.",
        "For cohort aggregate analysis, use the confidential revocation form at ? by ?. For an individual opt-in record, contact the independent CARE staff member at ? before the operational linkage and submission data are destroyed one month after the quiz deadline. After a record has been irreversibly de-identified and combined into non-identifiable results, it may no longer be possible to locate and remove it.":
            "If you opt in and later change your mind, contact the independent CARE staff member at ? by ? (no later than one month after the quiz deadline). After the linkage is destroyed or the record has been irreversibly de-identified and incorporated into non-identifiable results, it may no longer be possible to locate and remove it.",
        "If you explicitly opt in to the individual five-year research record, the de-identified data may be used in future educational research that has appropriate ethics approval. Cohort aggregate use and the anonymous survey do not imply consent to retain a linked individual record.":
            "If you explicitly opt in, the de-identified record may be used in individual and aggregate analyses for this study and future educational research that has appropriate ethics approval. If you do not opt in, your RepoProof record will not be retained or included in research analyses.",
    }
    replace_exact(doc, mapping, "PIS")
    doc.save(path)
    return path

def update_consent():
    path = DONE / "RepoProof_COMP5339_Individual_Record_EConsent_DRAFT.docx"
    doc = Document(path)
    mapping = {
        "COMP5339 students — optional individual five-year research record":
            "COMP5339 students — optional five-year RepoProof research record",
        "This consent is only for retaining and analysing my individual de-identified RepoProof research record for five years. It is separate from completing the quiz, the cohort aggregate opt-out pathway, and the optional anonymous survey. In giving consent, I confirm that:":
            "I voluntarily consent to my de-identified RepoProof research record being retained for five years and used in individual and aggregate research analyses. This choice is separate from completing the quiz and from the optional anonymous survey. In giving consent, I confirm that:",
        "I consent to my de-identified individual quiz record being retained for five years and used for this study and future educational research only where appropriate ethics approval is in place.":
            "I consent to my de-identified quiz and code-derived research record being retained for five years and used in individual and aggregate analyses for this study and future educational research only where appropriate ethics approval is in place.",
        "I understand that results may be published only in de-identified aggregate form and will not identify me.":
            "I understand that analyses include consenting participants only, participation rates will be reported, and results will not be described as necessarily representative of the full COMP5339 cohort.",
        "☐ YES — I consent to five-year retention and research use of my de-identified individual RepoProof record.":
            "☐ YES — I consent to five-year retention and individual and aggregate research use of my de-identified RepoProof record.",
        "☐ NO — do not retain my individual RepoProof record for research. This does not affect quiz credit or marks.":
            "☐ NO — do not retain or use my RepoProof record for research, including aggregate research analysis. This does not affect quiz access, feedback, credit or marks.",
    }
    replace_exact(doc, mapping, "Consent")
    doc.save(path)
    return path

def update_recruitment():
    src = OLD / "RepoProof_COMP5339_Recruitment_OptIn_Individual_Record_and_Survey_DRAFT.docx"
    out = DONE / "RepoProof_COMP5339_Recruitment_OptIn_All_RepoProof_Research_Use_and_Survey_DRAFT.docx"
    doc = Document(src)
    mapping = {
        "(Optional individual five-year research record and anonymous experience survey)":
            "(Optional five-year RepoProof research record and anonymous experience survey)",
        "We invite you to choose whether your individual de-identified RepoProof quiz record may be retained for five years for educational research. We also invite you to complete a separate anonymous experience survey.":
            "We invite you to choose whether your de-identified RepoProof quiz and code-derived research record may be retained for five years and used in individual and aggregate educational research analyses. We also invite you to complete a separate anonymous experience survey.",
        "Review the PIS and e-consent at ?. Select YES only if you consent to five-year retention of your de-identified individual record. Selecting NO, or not opting in, has no effect on the 2% quiz-completion credit or any other mark. The optional anonymous survey is at ?.":
            "Review the PIS and e-consent at ?. Select YES only if you consent to five-year retention and all research use of your de-identified RepoProof record, including aggregate analysis. If you select NO or do not opt in, your record will not be retained or used for research. This has no effect on the 2% quiz-completion credit or any other mark. The optional anonymous survey is at ?.",
    }
    replace_exact(doc, mapping, "Recruitment")
    doc.save(out)
    return out

if __name__ == "__main__":
    for p in (update_pis(), update_consent(), update_recruitment()):
        print(p)
